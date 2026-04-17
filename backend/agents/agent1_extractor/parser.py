"""
parser.py — Agent 1: CV Extractor
=================================
Entry point: parse_cv(source, job_description=None)

Supported input types
---------------------
- str path ending in .pdf  → LlamaParse (cloud), PyMuPDF fallback
- str path ending in .docx → python-docx
- any other str            → treated as raw CV text (plain-text pass-through)

Output
------
A dict matching the CVExtraction dataclass schema defined in schema.py.
This dict is the contract consumed by Agent 2 (OSINT) and Agent 3 (Verifier).
"""

from __future__ import annotations

import json
import logging
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from dotenv import load_dotenv

# ── Bootstrap ─────────────────────────────────────────────────────────────────

load_dotenv()

logger = logging.getLogger(__name__)

# ── Internal helpers: Text extraction ─────────────────────────────────────────

def _extract_pdf_llamaparse(pdf_path: str) -> tuple[str, str]:
    """
    Extract text from a PDF using LlamaParse (cloud).
    Returns (text, parse_method).
    Falls back to PyMuPDF on any error and records the fallback in the returned method tag.
    """
    import llama_parse  # llama-parse package

    api_key = os.getenv("LLAMA_CLOUD_API_KEY")
    if not api_key:
        raise EnvironmentError("LLAMA_CLOUD_API_KEY is not set in environment.")

    try:
        parser = llama_parse.LlamaParse(
            api_key=api_key,
            result_type="markdown",
            system_prompt="Extract all text and preserve all hyperlinks, especially LinkedIn and GitHub.", 
            verbose=False,
        )
        documents = parser.load_data(pdf_path)
        text = "\n\n".join(doc.text for doc in documents)
        if text.strip():
            return text, "llamaparse"
        
        # LlamaParse returned empty — try fallback
        logger.warning("LlamaParse returned empty text for %s; falling back to PyMuPDF.", pdf_path)
    except Exception as exc:
        logger.warning("LlamaParse failed for %s (%s); falling back to PyMuPDF.", pdf_path, exc)

    return _extract_pdf_pymupdf(pdf_path)


def _extract_pdf_pymupdf(pdf_path: str) -> tuple[str, str]:
    """Fallback PDF extraction using local PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        final_content = ""
        all_links = []

        for page in doc:
            # 1. Lấy text thuần
            page_text = page.get_text()
            
            # 2. Lấy link ẩn và đính kèm trực tiếp vào text của trang đó
            links = page.get_links()
            page_links_text = ""
            for link in links:
                if "uri" in link:
                    uri = link["uri"]
                    all_links.append(uri)
                    # Tạo một dòng chú thích để Gemini biết link này tồn tại
                    page_links_text += f"\n[Hyperlink found: {uri}]\n"
            
            final_content += page_text + page_links_text

        doc.close()

        # Thêm một section tổng hợp ở cuối để chắc chắn Agent 1 không sót
        if all_links:
            final_content += "\n\n--- EXTRACTED HYPERLINKS (METADATA) ---\n"
            final_content += "\n".join(list(set(all_links)))

        return final_content, "pymupdf_enhanced"
    except Exception as e:
        return str(e), "error"


def _extract_docx(docx_path: str) -> tuple[str, str]:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document  # python-docx

        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also pull text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs), "docx"
    except ImportError:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")


def extract_text(source: str) -> tuple[str, str, str]:
    """
    Detect input type and return (raw_text, parse_method, source_label).

    Args:
        source: A file path (.pdf / .docx) or a raw CV text string.

    Returns:
        raw_text:     Plain text content of the CV.
        parse_method: One of "llamaparse", "pymupdf_fallback", "docx", "plaintext".
        source_label: File name (for metadata) or "plaintext".
    """
    lower = source.lower()

    if lower.endswith(".pdf") and Path(source).exists():
        text, method = _extract_pdf_llamaparse(source)
        return text, method, Path(source).name

    if lower.endswith(".docx") and Path(source).exists():
        text, method = _extract_docx(source)
        return text, method, Path(source).name

    # Treat as raw text
    return source, "plaintext", "plaintext"


# ── Internal helper: Gemini call ──────────────────────────────────────────────

def _call_gemini(system_prompt: str, user_prompt: str, model: str = "gemini-2.5-flash") -> str:
    """Call Gemini and return the raw response string."""
    from google import genai
    from google.genai import types

    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise EnvironmentError("GOOGLE_API_KEY is not set in environment.")

    client = genai.Client(api_key=api_key)

    response = client.models.generate_content(
        model=model,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=system_prompt,
            response_mime_type="application/json",
            temperature=0.0,   # deterministic extraction — no creativity needed
        ),
    )
    return response.text


# ── Public API ────────────────────────────────────────────────────────────────

def parse_cv(
    source: str,
    job_description: str | None = None,
    gemini_model: str = "gemini-2.5-flash",
) -> dict[str, Any]:
    """
    Main entry point for Agent 1.

    Args:
        source:          File path to a PDF/DOCX, or raw CV text string.
        job_description: Optional job description to guide relevance tagging.
        gemini_model:    Gemini model name. Override for testing with a cheaper model.

    Returns:
        A dict matching the CVExtraction schema. Ready for Agent 2 and Agent 3.

    Raises:
        ValueError:       If Gemini returns malformed JSON after a retry.
        EnvironmentError: If required API keys are missing.
    """
    # Import theo cấu trúc module nội bộ
    from .prompt import SYSTEM_PROMPT, build_retry_prompt, build_user_prompt
    from .schema import CVExtraction

    warnings: list[str] = []

    # ---- Step 1: extract raw text ----------------------------------------
    raw_text, parse_method, source_label = extract_text(source)

    if parse_method == "pymupdf_fallback":
        warnings.append("LlamaParse unavailable or returned empty; used PyMuPDF as fallback.")

    if not raw_text.strip():
        raise ValueError(f"Could not extract any text from source: {source!r}")

    # ---- Step 2: build prompt --------------------------------------------
    user_prompt = build_user_prompt(raw_text, job_description)

    # ---- Step 3: call Gemini (with one retry on JSON parse failure) ------
    raw_response: str = ""
    try:
        raw_response = _call_gemini(SYSTEM_PROMPT, user_prompt, model=gemini_model)
        extracted_dict = json.loads(raw_response)

    except json.JSONDecodeError as first_err:
        logger.warning(
            "Gemini returned malformed JSON on first attempt (%s). Retrying...", first_err
        )
        warnings.append(f"First extraction attempt returned invalid JSON: {first_err}. Retry succeeded.")

        retry_prompt = build_retry_prompt(raw_text, raw_response, str(first_err))
        try:
            raw_response = _call_gemini(SYSTEM_PROMPT, retry_prompt, model=gemini_model)
            extracted_dict = json.loads(raw_response)
        except json.JSONDecodeError as second_err:
            raise ValueError(
                f"Gemini returned malformed JSON twice. Last error: {second_err}\n"
                f"Last response snippet: {raw_response[:500]}"
            ) from second_err

    # ---- Step 4: inject metadata ----------------------------------------
    extracted_dict["metadata"] = {
        "source_file": source_label,
        "parse_method": parse_method,
        "extraction_model": gemini_model,
        "extracted_at": datetime.now(timezone.utc).isoformat(),
        "warnings": warnings,
    }

    # ---- Step 5: validate with Dataclass (thay cho Pydantic) -------------
    try:
        cv_obj = CVExtraction.from_dict(extracted_dict)
    except Exception as validation_err:
        # Xảy ra khi mô hình sinh cấu trúc bị sai sót trầm trọng
        logger.error("Schema validation/mapping failed: %s", validation_err)
        extracted_dict["metadata"]["warnings"].append(
            f"Schema mapping warning: {validation_err}. Output may contain type mismatches."
        )
        return extracted_dict

    result = cv_obj.to_dict()

    # ---- Step 6: attach convenience views for downstream agents ----------
    result["_osint_targets"] = cv_obj.osint_targets()
    result["_all_urls"] = cv_obj.all_urls()

    return result


# ── CLI convenience ───────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python -m backend.agents.agent1_extractor.parser <cv_file_or_text> [job_description]")
        sys.exit(1)

    source_arg = sys.argv[1]
    jd_arg = sys.argv[2] if len(sys.argv) > 2 else None

    output = parse_cv(source_arg, job_description=jd_arg)
    print(json.dumps(output, indent=2, ensure_ascii=False))